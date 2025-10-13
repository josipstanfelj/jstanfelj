import customtkinter as ctk
from tkinter import filedialog, Tk
import os
from docx import Document
from fpdf import FPDF
from PyPDF2 import PdfReader, PdfWriter
import tkinter as tk


class DocumentConverterApp:
    @staticmethod
    def make_screen(frame):
        for widget in frame.winfo_children():
            widget.destroy()

       
        label = ctk.CTkLabel(frame, 
                             text="Document Converter", 
                             font=("Arial", 30), 
                             text_color="#dfe6e9")
        label.grid(row=0, column=0, columnspan=3, sticky="nsew", pady=(20, 10))
        
        select_button = ctk.CTkButton(frame, 
                                      text="Upload Document", 
                                      font=("Arial", 20), 
                                      command=lambda: DocumentConverterApp.select_document(frame))
        select_button.grid(row=2, column=0, columnspan=3, sticky="nsew", pady=(10, 10))

       
        format_label = ctk.CTkLabel(frame, 
                                    text="Select Output Format:", 
                                    font=("Arial", 20))
        format_label.grid(row=3, column=0, columnspan=1, sticky="e", padx=(10, 5), pady=(10, 10))

      
        format_var = ctk.StringVar(value="PDF")
        format_dropdown = ctk.CTkComboBox(frame, 
                                          values=["PDF", "DOCX", "TXT", "RTF"], 
                                          variable=format_var, 
                                          font=("Arial", 20))
        format_dropdown.grid(row=3, column=1, columnspan=2, sticky="w", padx=(5, 10), pady=(10, 10))

       
        convert_button = ctk.CTkButton(frame, 
                                       text="Convert", 
                                       font=("Arial", 20), 
                                       command=lambda: DocumentConverterApp.convert_document(format_var.get(), frame))
        convert_button.grid(row=4, column=0, columnspan=3, sticky="nsew", pady=(10, 10))

        
        DocumentConverterApp.result_label = ctk.CTkLabel(frame, 
                                                         text="", 
                                                         font=("Arial", 20), 
                                                         text_color="#dfe6e9")
        DocumentConverterApp.result_label.grid(row=5, column=0, columnspan=3, sticky="nsew", pady=(10, 10))

        
        DocumentConverterApp.download_button = ctk.CTkButton(frame, 
                                                              text="Download Converted File", 
                                                              font=("Arial", 20), 
                                                              state="disabled", 
                                                              command=lambda: DocumentConverterApp.download_converted_file())
        DocumentConverterApp.download_button.grid(row=6, column=0, columnspan=3, sticky="nsew", pady=(10, 10))

        
        frame.grid_columnconfigure(0, weight=1)
        frame.grid_columnconfigure(1, weight=1)
        frame.grid_columnconfigure(2, weight=1)

    selected_file = ""
    converted_file = ""

    @staticmethod
    def on_drag_enter(event):
        event.widget.configure(text="Drop File to Upload")

    @staticmethod
    def on_file_drop(event, frame):
        file = event.data  
        if os.path.isfile(file):
            DocumentConverterApp.selected_file = file
            DocumentConverterApp.result_label.configure(
                text=f"Selected file: {os.path.basename(file)}"
            )
        else:
            DocumentConverterApp.result_label.configure(
                text="Invalid file. Please try again."
            )

    @staticmethod
    def select_document(frame):
        
        file = filedialog.askopenfilename(
            filetypes=[("Document files", "*.docx;*.pdf;*.txt;*.rtf")]
        )

       
        DocumentConverterApp.selected_file = file

        
        if file:
            DocumentConverterApp.result_label.configure(
                text=f"Selected file: {os.path.basename(file)}"
            )
        else:
            DocumentConverterApp.result_label.configure(
                text="No document selected."
            )



    @staticmethod
    def get_unique_filename(directory, filename):
        base, ext = os.path.splitext(filename)
        counter = 1
        new_filename = filename
        while os.path.exists(os.path.join(directory, new_filename)):
            new_filename = f"{base}({counter}){ext}"
            counter += 1
        return new_filename

    @staticmethod
    def convert_document(output_format, frame):

        current_dir = os.path.dirname(os.path.abspath(__file__))
        font_path = os.path.join(current_dir, "DejaVuSans.ttf")

        if not DocumentConverterApp.selected_file:
            DocumentConverterApp.result_label.configure(
                text="Please upload a document first."
            )
            return

        input_file = DocumentConverterApp.selected_file
        base_name = os.path.splitext(os.path.basename(input_file))[0]

        try:
            if output_format == "PDF" and input_file.endswith(".docx"):
                doc = Document(input_file)
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                pdf.set_font("Arial", size=12)

                pdf.add_font("DejaVu", "", font_path, uni=True)
                pdf.set_font("DejaVu", size=12)

                for paragraph in doc.paragraphs:
                    pdf.multi_cell(0, 10, paragraph.text)

                output_file = f"{base_name}.pdf"
                pdf.output(output_file)

            elif output_format == "DOCX" and input_file.endswith(".pdf"):
                reader = PdfReader(input_file)
                doc = Document()

                for page in reader.pages:
                    doc.add_paragraph(page.extract_text())

                output_file = f"{base_name}.docx"
                doc.save(output_file)

            elif output_format == "TXT" and input_file.endswith(".docx"):
                doc = Document(input_file)
                output_file = f"{base_name}.txt"
                with open(output_file, "w", encoding="utf-8") as txt_file:
                    for paragraph in doc.paragraphs:
                        txt_file.write(paragraph.text + "\n")

            elif output_format == "DOCX" and input_file.endswith(".txt"):
                with open(input_file, "r", encoding="utf-8") as txt_file:
                    content = txt_file.readlines()
                doc = Document()
                for line in content:
                    doc.add_paragraph(line.strip())
                output_file = f"{base_name}.docx"
                doc.save(output_file)

            elif output_format == "PDF" and input_file.endswith(".txt"):
                with open(input_file, "r", encoding="utf-8") as txt_file:
                    content = txt_file.readlines()
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                
                pdf.add_font("DejaVu", "", font_path, uni=True)
                pdf.set_font("DejaVu", size=12)
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                for line in content:
                    pdf.multi_cell(0, 10, line.strip())
                output_file = f"{base_name}.pdf"
                pdf.output(output_file)

            elif output_format == "TXT" and input_file.endswith(".pdf"):
                reader = PdfReader(input_file)
                output_file = f"{base_name}.txt"
                with open(output_file, "w", encoding="utf-8") as txt_file:
                    for page in reader.pages:
                        txt_file.write(page.extract_text() + "\n")

            elif output_format == "RTF" and input_file.endswith(".docx"):
                doc = Document(input_file)
                output_file = f"{base_name}.rtf"
                with open(output_file, "w", encoding="utf-8") as rtf_file:
                    for paragraph in doc.paragraphs:
                        rtf_file.write(paragraph.text + "\n")

            elif output_format == "DOCX" and input_file.endswith(".rtf"):
                with open(input_file, "r", encoding="utf-8") as rtf_file:
                    try:
                        content = rtf_file.readlines()
                    except UnicodeDecodeError:
                        DocumentConverterApp.result_label.configure(
                            text="Error: RTF file contains invalid characters."
                        )
                        return
                doc = Document()
                for line in content:
                    doc.add_paragraph(line.strip())
                output_file = f"{base_name}.docx"
                doc.save(output_file)

            elif output_format == "PDF" and input_file.endswith(".rtf"):
                with open(input_file, "r", encoding="utf-8") as rtf_file:
                    try:
                        content = rtf_file.readlines()
                    except UnicodeDecodeError:
                        DocumentConverterApp.result_label.configure(
                            text="Error: RTF file contains invalid characters."
                        )
                        return
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                pdf.add_font("DejaVu", "", font_path, uni=True)
                pdf.set_font("DejaVu", size=12)
                for line in content:
                    pdf.multi_cell(0, 10, line.strip())
                output_file = f"{base_name}.pdf"
                pdf.output(output_file)

            DocumentConverterApp.converted_file = output_file
            DocumentConverterApp.result_label.configure(
                text="Conversion successful! Ready to download."
            )
            DocumentConverterApp.download_button.configure(state="normal")

        except Exception as e:
            DocumentConverterApp.result_label.configure(
                text=f"Error during conversion: {str(e)}"
            )

    @staticmethod
    def download_converted_file():
        if not DocumentConverterApp.converted_file:
            DocumentConverterApp.result_label.configure(
                text="No file to download."
            )
            return

        downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
        if not os.path.exists(downloads_folder):
            os.makedirs(downloads_folder)

        unique_filename = DocumentConverterApp.get_unique_filename(
            downloads_folder, os.path.basename(DocumentConverterApp.converted_file)
        )
        destination = os.path.join(downloads_folder, unique_filename)
        os.rename(DocumentConverterApp.converted_file, destination)

        DocumentConverterApp.result_label.configure(
            text=f"File has been saved to Downloads as {unique_filename}."
        )
